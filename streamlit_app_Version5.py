import streamlit as st
import streamlit.components.v1 as components
from docx import Document
import re
import uuid
import os
import time
import html
import csv
from io import BytesIO

# ----------------------------------------------------
# إعدادات الصفحة الأساسية
# ----------------------------------------------------
st.set_page_config(
    page_title="القوانين اليمنية بآخر تعديلاتها حتى عام 2025م",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ----------------- دعم الوضع الليلي -----------------
if "night_mode" not in st.session_state:
    st.session_state.night_mode = False

# ----------------------------------------------------
# دعم الكتابة من اليمين
st.markdown("""
<style>
textarea, input[type="text"], .stTextArea textarea, .stTextInput input {
    direction: rtl !important;
    text-align: right !important;
    font-family: "Tahoma", "Arial", sans-serif !important;
    font-size: 18px !important;
}
[data-testid="stTextArea"] textarea,
[data-testid="stTextInput"] input {
    direction: rtl !important;
    text-align: right !important;
}
</style>
""", unsafe_allow_html=True)

# ----------------------------------------------------
# بقية المتغيرات
TRIAL_DURATION = 3 * 24 * 60 * 60  # 3 أيام
TRIAL_USERS_FILE = "trial_users.txt"
DEVICE_ID_FILE = "device_id.txt"
ACTIVATED_FILE = "activated.txt"
ACTIVATION_CODES_FILE = "activation_codes.txt"
LAWS_DIR = "laws"

def get_device_id():
    if os.path.exists(DEVICE_ID_FILE):
        with open(DEVICE_ID_FILE, "r") as f:
            return f.read().strip()
    new_id = str(uuid.uuid4())
    with open(DEVICE_ID_FILE, "w") as f:
        f.write(new_id)
    return new_id

def get_trial_start(device_id):
    if not os.path.exists(TRIAL_USERS_FILE):
        return None
    with open(TRIAL_USERS_FILE, "r") as f:
        reader = csv.reader(f)
        for row in reader:
            if row and row[0] == device_id:
                return float(row[1])
    return None

def register_trial(device_id):
    if not os.path.exists(TRIAL_USERS_FILE):
        with open(TRIAL_USERS_FILE, "w", newline='') as f:
            pass
    with open(TRIAL_USERS_FILE, "a", newline='') as f:
        writer = csv.writer(f)
        writer.writerow([device_id, time.time()])

def is_activated():
    return os.path.exists(ACTIVATED_FILE)

def activate_app(code):
    if not os.path.exists(ACTIVATION_CODES_FILE):
        return False
    with open(ACTIVATION_CODES_FILE, "r") as f:
        codes = [line.strip() for line in f.readlines()]
    if code in codes:
        codes.remove(code)
        with open(ACTIVATION_CODES_FILE, "w") as f:
            for c in codes:
                f.write(c + "\n")
        with open(ACTIVATED_FILE, "w") as f:
            f.write("activated")
        return True
    return False

def highlight_keywords(text, keywords, exact_match=False):
    # تمييز الكلمات المفتاحية في النص الأصلي (دون تبسيط)
    for kw in keywords:
        if not kw:
            continue
        if exact_match:
            # تطابق تام: ابحث عن الكلمة منفردة فقط
            pattern = re.compile(r"(?<!\w)"+re.escape(kw)+r"(?!\w)", re.IGNORECASE)
        else:
            # بحث عادي: أي تطابق جزئي
            pattern = re.compile(re.escape(kw), re.IGNORECASE)
        text = pattern.sub(f'<mark>\\g<0></mark>', text)
    return text

def export_results_to_word(results, filename="نتائج_البحث.docx"):
    from docx import Document
    document = Document()
    document.add_heading('نتائج البحث في القوانين اليمنية', level=1)
    
    if not results:
        document.add_paragraph("لم يتم العثور على نتائج للكلمات المفتاحية المحددة.")
    else:
        for i, r in enumerate(results):
            document.add_heading(f"القانون: {r['law']} - المادة: {r['num']}", level=2)
            document.add_paragraph(r['plain'])
            if i < len(results) - 1:
                document.add_page_break() 

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

def normalize_arabic_numbers(text):
    arabic_to_english = str.maketrans('٠١٢٣٤٥٦٧٨٩', '0123456789')
    return text.translate(arabic_to_english)

# ----------------- دالة تبسيط النص العربي المتقدمة -----------------
def normalize_arabic_text(text):
    # إزالة المد (الحروف المكررة، خاصة الألف أو أي حرف آخر أكثر من مرتين)
    text = re.sub(r'(.)\1{2,}', r'\1', text)
    # إزالة التشكيل والتنوين والحركات
    text = re.sub(r'[\u064B-\u0652]', '', text)
    # توحيد الألف والهمزات
    text = re.sub('[إأآا]', 'ا', text)
    # توحيد الياء والألف المقصورة
    text = re.sub('[ىي]', 'ي', text)
    # توحيد التاء المربوطة مع الهاء
    text = re.sub('[ة]', 'ه', text)
    # توحيد الواو الصغيرة
    text = re.sub('ؤ', 'و', text)
    # توحيد همزة على نبرة
    text = re.sub('ئ', 'ي', text)
    # إزالة علامات الترقيم
    text = re.sub(r'[^\w\s]', '', text)
    # إزالة المسافات الزائدة
    text = re.sub('\s+', ' ', text)
    return text.strip()

def run_main_app():
    # ----------- الوضع الليلي -----------
    with st.sidebar:
        st.session_state.night_mode = st.toggle("🌙 تفعيل الوضع الليلي", value=st.session_state.night_mode)

    # CSS للوضع الليلي والنهاري + نتائج البحث
    if st.session_state.night_mode:
        st.markdown("""
        <style>
        body, .stApp {
            background-color: #181a1b !important;
            color: #f1f1f1 !important;
        }
        .stTextInput input, .stTextArea textarea, textarea, input[type="text"] {
            background-color: #222426 !important;
            color: #f1f1f1 !important;
        }
        .stButton button, .stDownloadButton button {
            background: linear-gradient(90deg, #333 0%, #222 100%) !important;
            color: #f1f1f1 !important;
            border: 1px solid #444 !important;
        }
        .stExpanderHeader, .stForm, .stMetric {
            background-color: #232526 !important;
            color: #f1f1f1 !important;
        }
        mark {
            background: #405983 !important;
            color: #fff !important;
        }
        .copy-material-btn {
            background: linear-gradient(90deg, #384e5a 0%, #213b4b 100%) !important;
            color: #eee !important;
        }
        .copy-material-btn:hover {
            background: linear-gradient(90deg, #213b4b 0%, #384e5a 100%) !important;
        }
        .result-box-night {
            background-color: #232526 !important;
            color: #fafafa !important;
            padding: 20px;
            margin-bottom: 10px;
            width: 100%;
            max-width: 100%;
            border-radius: 10px;
            border: 1px solid #333;
            direction: rtl;
            text-align: right;
        }
        </style>
        """, unsafe_allow_html=True)
    else:
        st.markdown("""
        <style>
        body, .stApp {
            background-color: #fff !important;
            color: #232323 !important;
        }
        .result-box-night {
            background-color: #f1f8e9 !important;
            color: #232323 !important;
            padding: 20px;
            margin-bottom: 10px;
            width: 100%;
            max-width: 100%;
            border-radius: 10px;
            border: 1px solid #c5e1a5;
            direction: rtl;
            text-align: right;
        }
        </style>
        """, unsafe_allow_html=True)

    # أزرار التمرير
    components.html("""
    <style>
    .scroll-btn {
        position: fixed;
        left: 10px;
        padding: 12px;
        font-size: 24px;
        border-radius: 50%;
        background-color: #c5e1a5;
        color: black;
        cursor: pointer;
        z-index: 9999;
        border: none;
        box-shadow: 1px 1px 5px #888;
    }
    #scroll-top-btn { bottom: 80px; }
    #scroll-bottom-btn { bottom: 20px; }
    .rtl-metric {
        direction: rtl;
        text-align: right !important;
        margin-right: 0 !important;
    }
    .rtl-metric .stMetric {
        text-align: right !important;
        direction: rtl;
    }
    .rtl-metric .stMetricDelta {
        display: block !important;
        text-align: right !important;
        direction: rtl;
    }
    .rtl-download-btn {
        direction: rtl;
        text-align: right !important;
        margin-right: 0 !important;
        display: flex;
        flex-direction: row-reverse;
        justify-content: flex-start;
    }
    textarea, .stTextArea, .stTextArea textarea, input[type="text"], .stTextInput input, .stTextInput textarea {
        direction: rtl !important;
        text-align: right !important;
        padding-right: 10px;
        font-family: "Tahoma", "Arial", sans-serif;
        font-size: 16px;
        line-height: 1.5;
    }
    .stButton, .stDownloadButton, .stMetric {
        direction: rtl !important;
        text-align: right !important;
    }
    </style>
    <button class='scroll-btn' id='scroll-top-btn' onclick='window.scrollTo({top: 0, behavior: "smooth"});'>⬆️</button>
    <button class='scroll-btn' id='scroll-bottom-btn' onclick='window.scrollTo({top: document.body.scrollHeight, behavior: "smooth"});'>⬇️</button>
    """, height=1)

    if not os.path.exists(LAWS_DIR):
        st.error(f"⚠️ مجلد '{LAWS_DIR}/' غير موجود. يرجى التأكد من وجود ملفات القوانين.")
        return

    files = [f for f in os.listdir(LAWS_DIR) if f.endswith(".docx")]
    if not files:
        st.warning(f"📂 لا توجد ملفات قوانين في مجلد '{LAWS_DIR}/'.")
        return

    st.markdown("""
        <div style="direction: rtl; text-align: right;">
        <h3 style="display: flex; align-items: center; gap: 10px;">🔎 نموذج البحث</h3>
        </div>
    """, unsafe_allow_html=True)
    with st.form("main_search_form"):
        st.markdown('<div style="direction: rtl; text-align: right;">اختر قانونًا للبحث:</div>', unsafe_allow_html=True)
        selected_file_form = st.selectbox("", ["الكل"] + files, key="main_file_select", label_visibility="collapsed")
        st.markdown('<div style="direction: rtl; text-align: right;">📌 اكتب كلمة أو جملة للبحث عنها:</div>', unsafe_allow_html=True)
        st.markdown('<div dir="rtl">', unsafe_allow_html=True)
        keywords_form = st.text_area(
            "",
            key="main_keywords_input",
            help="أدخل الكلمات التي تريد البحث عنها، وافصل بينها بفاصلة إذا كانت أكثر من كلمة.",
        )
        st.markdown('</div>', unsafe_allow_html=True)
        st.markdown('<div style="direction: rtl; text-align: right;">أو أبحث برقم المادة:</div>', unsafe_allow_html=True)
        st.markdown('<div dir="rtl">', unsafe_allow_html=True)
        article_number_input = st.text_input(
            "",
            key="article_number_input",
            help="أدخل رقم المادة للبحث عنها مباشرة (يمكن استخدام أرقام عربية أو إنجليزية)."
        )
        st.markdown('</div>', unsafe_allow_html=True)
        # خيار البحث المتقدم: تطابق تام أو بحث مرن
        advanced_search_col = st.columns([1, 2, 5])
        with advanced_search_col[2]:
            exact_match = st.checkbox("تطابق تام للكلمة (لا تظهر مشتقاتها مثل تظلم عند البحث عن ظلم)", key="exact_match_checkbox")
        search_btn_col = st.columns([1, 2, 12])
        with search_btn_col[2]:
            submitted = st.form_submit_button("🔍 بدء البحث", use_container_width=True)

    if "results" not in st.session_state:
        st.session_state.results = []
    if "search_done" not in st.session_state:
        st.session_state.search_done = False

    if submitted:
        results = []
        search_files = files if selected_file_form == "الكل" else [selected_file_form]
        kw_list = [k.strip() for k in keywords_form.split(",") if k.strip()] if keywords_form else []
        search_by_article = bool(article_number_input.strip())

        # تبسيط كلمات البحث
        normalized_kw_list = [normalize_arabic_text(kw) for kw in kw_list] if kw_list else []
        norm_article = normalize_arabic_numbers(article_number_input.strip()) if search_by_article else ""

        with st.spinner("جاري البحث في القوانين... قد يستغرق الأمر بعض الوقت."):
            for file in search_files:
                try:
                    doc = Document(os.path.join(LAWS_DIR, file))
                except Exception as e:
                    st.warning(f"⚠️ تعذر قراءة الملف {file}: {e}. يرجى التأكد من أنه ملف DOCX صالح.")
                    continue

                law_name = file.replace(".docx", "")
                last_article = "غير معروفة"
                current_article_paragraphs = []

                for para in doc.paragraphs:
                    txt = para.text.strip()
                    if not txt:
                        continue
                    match = re.match(r"مادة\s*[\(]?\s*(\d+)[\)]?", txt)
                    if match:
                        if current_article_paragraphs:
                            full_text = "\n".join(current_article_paragraphs)
                            add_result = False

                            # --- بحث مرن للكلمات أو تطابق تام حسب اختيار المستخدم ---
                            simple_full_text = normalize_arabic_text(full_text)
                            if search_by_article and normalize_arabic_numbers(last_article) == norm_article:
                                add_result = True
                            elif normalized_kw_list:
                                for idx, kw in enumerate(normalized_kw_list):
                                    if not kw:
                                        continue
                                    if exact_match:
                                        # تطابق تام: الكلمة منفصلة فقط
                                        # استخدم حدود الكلمة
                                        pattern = r'(?<!\w)'+re.escape(kw)+r'(?!\w)'
                                        if re.search(pattern, simple_full_text):
                                            add_result = True
                                            break
                                    else:
                                        # بحث مرن: أي تطابق جزئي
                                        if kw in simple_full_text:
                                            add_result = True
                                            break

                            if add_result:
                                highlighted = highlight_keywords(full_text, kw_list, exact_match=exact_match) if kw_list else full_text
                                results.append({
                                    "law": law_name,
                                    "num": last_article,
                                    "text": highlighted,
                                    "plain": full_text
                                })
                            current_article_paragraphs = []
                        last_article = match.group(1)
                    current_article_paragraphs.append(txt)

                if current_article_paragraphs:
                    full_text = "\n".join(current_article_paragraphs)
                    add_result = False
                    simple_full_text = normalize_arabic_text(full_text)
                    if search_by_article and normalize_arabic_numbers(last_article) == norm_article:
                        add_result = True
                    elif normalized_kw_list:
                        for idx, kw in enumerate(normalized_kw_list):
                            if not kw:
                                continue
                            if exact_match:
                                pattern = r'(?<!\w)'+re.escape(kw)+r'(?!\w)'
                                if re.search(pattern, simple_full_text):
                                    add_result = True
                                    break
                            else:
                                if kw in simple_full_text:
                                    add_result = True
                                    break

                    if add_result:
                        highlighted = highlight_keywords(full_text, kw_list, exact_match=exact_match) if kw_list else full_text
                        results.append({
                            "law": law_name,
                            "num": last_article,
                            "text": highlighted,
                            "plain": full_text
                        })

        st.session_state.results = results
        st.session_state.search_done = True
        if not results:
            st.info("لم يتم العثور على نتائج مطابقة للبحث.")

    if st.session_state.get("search_done", False) and st.session_state.results:
        st.markdown("<h2 style='text-align: center; color: #388E3C;'>نتائج البحث في القوانين 📚</h2>", unsafe_allow_html=True)
        st.markdown("---")

    if st.session_state.get("search_done", False):
        results = st.session_state.results
        unique_laws = sorted(set(r["law"] for r in results))

        st.markdown('<div class="rtl-metric">', unsafe_allow_html=True)
        st.metric(label="📊 إجمالي النتائج التي تم العثور عليها", value=f"{len(results)}", delta=f"في {len(unique_laws)} قانون/ملف")
        st.markdown('</div>', unsafe_allow_html=True)

        if results:
            export_data = export_results_to_word(results)
            st.markdown('<div class="rtl-download-btn">', unsafe_allow_html=True)
            st.download_button(
                label="⬇️ تصدير النتائج إلى Word",
                data=export_data,
                file_name="نتائج_البحث_القوانين_اليمنية.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_button_word_main",
                use_container_width=False
            )
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.warning("لا توجد نتائج لتصديرها.")
            
        st.markdown("---")

        if results:
            st.markdown('<div style="direction: rtl; text-align: right;">فلترة النتائج حسب القانون:</div>', unsafe_allow_html=True)
            selected_law_filter = st.selectbox("", ["الكل"] + unique_laws, key="results_law_filter", label_visibility="collapsed")
            filtered = results if selected_law_filter == "الكل" else [r for r in results if r["law"] == selected_law_filter]

            for i, r in enumerate(filtered):
                with st.expander(f"📚 المادة ({r['num']}) من قانون {r['law']}", expanded=True):
                    st.markdown(f'''
                    <div class="result-box-night">
                        <p style="font-size:17px;line-height:1.8;margin-top:0px;">
                            {r["text"]}
                        </p>
                    </div>
                    ''', unsafe_allow_html=True)
                    components.html(f"""
                        <style>
                        .copy-material-btn {{
                            display: inline-flex;
                            align-items: center;
                            gap: 10px;
                            background: linear-gradient(90deg, #1abc9c 0%, #2980b9 100%);
                            color: #fff;
                            border: none;
                            border-radius: 30px;
                            font-size: 18px;
                            font-family: 'Cairo', 'Tajawal', sans-serif;
                            padding: 10px 22px;
                            cursor: pointer;
                            box-shadow: 0 4px 15px rgba(41, 128, 185, 0.4);
                            transition: all 0.3s ease;
                            margin-bottom: 10px;
                            direction: rtl;
                            white-space: nowrap;
                        }}
                        .copy-material-btn:hover {{
                            background: linear-gradient(90deg, #2980b9 0%, #1abc9c 100%);
                            box-shadow: 0 6px 20px rgba(41, 128, 185, 0.6);
                            transform: translateY(-2px);
                        }}
                        .copy-material-btn .copy-icon {{
                            font-size: 20px;
                            margin-left: 8px;
                            display: block;
                        }}
                        .copy-material-btn .copied-check {{
                            font-size: 20px;
                            color: #fff;
                            margin-left: 8px;
                            display: none;
                        }}
                        .copy-material-btn.copied .copy-icon {{
                            display: none;
                        }}
                        .copy-material-btn.copied .copied-check {{
                            display: inline;
                            animation: fadein-check 0.5s ease-out;
                        }}
                        @keyframes fadein-check {{
                            0% {{ opacity: 0; transform: scale(0.7); }}
                            100% {{ opacity: 1; transform: scale(1); }}
                        }}
                        </style>
                        <button class="copy-material-btn" id="copy_btn_{i}_{r['law']}_{r['num']}" onclick="
                            navigator.clipboard.writeText(document.getElementById('plain_text_{i}_{r['law']}_{r['num']}').innerText);
                            var btn = document.getElementById('copy_btn_{i}_{r['law']}_{r['num']}');
                            btn.classList.add('copied');
                            setTimeout(function(){{
                                btn.classList.remove('copied');
                            }}, 1800);
                        ">
                            <span class="copy-icon">
                                <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
                                    <rect x="9" y="9" width="13" height="13" rx="2" ry="2"></rect>
                                    <path d="M5 15H4a2 2 0 0 1-2-2V4a2 2 0 0 1 2-2h9a2 2 0 0 1 2 2v1"></path>
                                </svg>
                            </span>
                            <span>نسخ</span>
                            <span class="copied-check">
                                <svg width="20" height="20" viewBox="0 0 24 24" fill="none" stroke="currentColor" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
                                    <polyline points="20 6 9 17 4 12"></polyline>
                                </svg>
                                تم النسخ!
                            </span>
                        </button>
                        <div id="plain_text_{i}_{r['law']}_{r['num']}" style="display:none;">{html.escape(r['plain'])}</div>
                    """, height=60)
        else:
            st.info("لا توجد نتائج لعرضها حاليًا. يرجى إجراء بحث جديد.")

def render_header():
    if os.path.exists("header.html"):
        with open("header.html", "r", encoding="utf-8") as f:
            header_html = f.read()
        st.markdown(header_html, unsafe_allow_html=True)
    else:
        st.error("⚠️ ملف 'header.html' غير موجود في مجلد المشروع.")

def main():
    render_header()

    device_id = get_device_id()
    trial_start = get_trial_start(device_id)

    if is_activated():
        run_main_app()
        return

    if trial_start is not None:
        elapsed_time = time.time() - trial_start
        remaining_time = int(TRIAL_DURATION - elapsed_time)
        if remaining_time > 0:
            run_main_app()
            return
        else:
            st.error("❌ انتهت مدة التجربة المجانية لهذا الجهاز. يرجى تفعيل التطبيق للاستمرار في الاستخدام.")

    st.markdown("""
    <div style='text-align:center; color:#2c3e50; font-size:22px; font-weight:bold; padding:20px;'>
        مرحباً بك عزيزي المستخدم، قم بالنقر على أيقونة بدء النسخة المجانية أو أدخل كود التفعيل:
    </div>""", unsafe_allow_html=True)

    with st.container(border=True):
        st.markdown("<h3 style='text-align:center; color:#2c3e50;'>⏱️ النسخة التجريبية المجانية</h3>", unsafe_allow_html=True)

        if trial_start is None:
            if st.button("🚀 بدء النسخة المجانية", key="start_trial_button", use_container_width=True):
                register_trial(device_id)
                st.rerun()

    st.markdown("---")

    with st.container(border=True):
        st.markdown("<h3 style='text-align:center; color:#2c3e50;'>🔐 النسخة المدفوعة</h3>", unsafe_allow_html=True)
        code = st.text_input("أدخل كود التفعيل هنا:", key="activation_code_input", help="أدخل الكود الذي حصلت عليه لتفعيل النسخة الكاملة.")
        if st.button("✅ تفعيل الآن", key="activate_button", use_container_width=True):
            if code and activate_app(code.strip()):
                st.success("✅ تم التفعيل بنجاح! يرجى إعادة تشغيل التطبيق لتطبيق التغييرات.")
                st.stop()
            else:
                st.error("❌ كود التفعيل غير صحيح أو انتهت صلاحيته.")

if __name__ == "__main__":
    main()